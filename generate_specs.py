from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ───────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=18, bold=True, color=(34, 85, 34))   # dark green
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # underline rule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '225522')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=(45, 110, 45))
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(60, 60, 60))
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{key}: ")
    set_font(r1, size=10.5, bold=True)
    r2 = p.add_run(value)
    set_font(r2, size=10.5)
    return p

def shade_row(row, hex_color="E8F5E9"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    hrow = t.rows[0]
    shade_row(hrow, "225522")
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_widths:
            cell.width = Cm(col_widths[i])
    # data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        if ri % 2 == 0:
            shade_row(tr, "F1F8F1")
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=10)
            if col_widths:
                cell.width = Cm(col_widths[ci])
    doc.add_paragraph()  # spacing after table


# ═══════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("VigorTerra")
set_font(run, size=36, bold=True, color=(34, 85, 34))

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Agricultural Yield Prediction Platform — Tunisia")
set_font(run2, size=15, italic=True, color=(80, 80, 80))

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("SOFTWARE SPECIFICATIONS DOCUMENT")
set_font(run3, size=13, bold=True, color=(45, 110, 45))

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Document title",   "VigorTerra — Software Specifications (Cahier de Charge)"),
    ("Version",          "1.0"),
    ("Date",             datetime.date.today().strftime("%B %d, %Y")),
    ("Project type",     "Web Application — Agricultural Decision Support System"),
    ("Domain",           "Smart Agriculture / Machine Learning / Data Engineering"),
    ("Target geography", "Tunisia"),
    ("Status",           "Draft — For Review"),
]
for k, v in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{k}:  ")
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(v)
    set_font(r2, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual)
# ═══════════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_entries = [
    ("1",   "Project Overview",                         3),
    ("2",   "Objectives and Scope",                     3),
    ("3",   "Stakeholders",                             4),
    ("4",   "Functional Requirements",                  4),
    ("4.1", "Authentication Module",                    4),
    ("4.2", "Home Dashboard",                           4),
    ("4.3", "Data Sources Page",                        5),
    ("4.4", "Dataset Overview Page",                    5),
    ("4.5", "ML Pipeline Page",                         5),
    ("4.6", "Yield Prediction (Test) Page",             5),
    ("4.7", "AI Agricultural Assistant (Ask Me)",       6),
    ("5",   "Non-Functional Requirements",              6),
    ("6",   "System Architecture",                      7),
    ("6.1", "Frontend Architecture",                    7),
    ("6.2", "Backend Architecture",                     7),
    ("6.3", "Data Flow",                                8),
    ("7",   "API Specification",                        8),
    ("8",   "Machine Learning Models",                  9),
    ("9",   "Data Sources",                             10),
    ("10",  "Technology Stack",                         10),
    ("11",  "Deployment",                               11),
    ("12",  "Constraints and Assumptions",              11),
    ("13",  "Glossary",                                 12),
]
for num, title, _ in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  {num}   {title}")
    set_font(r, size=10.5, bold=(len(num) == 1))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
h1("1. Project Overview")
body(
    "VigorTerra is a web-based decision-support platform designed to predict agricultural crop yield "
    "in Tunisia using climate and soil data. The platform combines machine learning models, an interactive "
    "React frontend, and a FastAPI backend to deliver actionable insights for farmers, agronomists, and "
    "agricultural researchers."
)
body(
    "The name 'VigorTerra' (vigour + terra / earth) reflects the project's goal: transforming raw "
    "environmental data into practical guidance that strengthens agricultural productivity."
)

h2("1.1  Context")
body(
    "Tunisian agriculture faces increasing pressure from climate variability, water scarcity, and "
    "soil degradation. Yield forecasting remains largely manual and reactive. VigorTerra bridges the "
    "gap between available open datasets (FAOSTAT, Agridata, Open-Meteo) and practical field decisions "
    "by providing a data-driven prediction layer accessible via a modern web interface."
)

h2("1.2  Key Features at a Glance")
features = [
    ("Yield Prediction", "Three ML models (Random Forest, Gradient Boosting, Stacking) predict yield in t/ha from 8 soil and climate inputs."),
    ("Anomaly Detection", "Automatic flagging when input parameters fall outside agronomic reference ranges."),
    ("AI Chat Assistant", "Terra — an AI agent powered by Google Gemini 1.5 Flash with a built-in agricultural knowledge base."),
    ("Multilingual UI", "Full support for English, French, and Arabic with dynamic RTL rendering."),
    ("Authentication", "Protected routes with login/sign-up/profile management."),
    ("Data Transparency", "Dedicated pages for data sources, dataset statistics, and ML pipeline visualization."),
]
add_table(["Feature", "Description"], features, col_widths=[4.5, 11])


# ═══════════════════════════════════════════════════════════════════════════
#  2. OBJECTIVES AND SCOPE
# ═══════════════════════════════════════════════════════════════════════════
h1("2. Objectives and Scope")

h2("2.1  Primary Objectives")
for obj in [
    "Provide an accessible, browser-based tool for yield prediction requiring no installation.",
    "Deliver interpretable results (yield quality label, anomaly reason) alongside raw predictions.",
    "Support evidence-based agricultural decisions through integrated data sources and AI guidance.",
    "Demonstrate a reproducible end-to-end ML pipeline from raw data collection to deployed API.",
]:
    bullet(obj)

h2("2.2  In Scope")
for item in [
    "React/Vite single-page application (frontend).",
    "FastAPI REST API with prediction, agent, and health endpoints (backend).",
    "Three approximation-based ML model implementations (Random Forest, Gradient Boosting, Stacking).",
    "AI chat assistant with local knowledge base and optional Gemini API fallback.",
    "Docker and docker-compose deployment configuration.",
    "Data collection scripts (scraping, FAOSTAT API, Open-Meteo API).",
    "Dataset preprocessing and merging pipeline (Python scripts).",
    "MLflow experiment notebooks (Jupyter).",
]:
    bullet(item)

h2("2.3  Out of Scope")
for item in [
    "Real-time IoT sensor integration.",
    "Mobile native application (iOS / Android).",
    "Payment or subscription management.",
    "Crop disease image recognition.",
    "Production-grade user database (current auth is client-side context only).",
]:
    bullet(item)


# ═══════════════════════════════════════════════════════════════════════════
#  3. STAKEHOLDERS
# ═══════════════════════════════════════════════════════════════════════════
h1("3. Stakeholders")
stakeholders = [
    ("Farmers / Agricultural Workers", "Primary end-users; enter field conditions to obtain yield forecasts."),
    ("Agronomists / Field Advisors",   "Use predictions and anomaly reports to guide crop management."),
    ("Agricultural Researchers",       "Explore dataset, pipeline, and model details for academic use."),
    ("Development Team",               "Engineers and data scientists building and maintaining the platform."),
    ("Ministry of Agriculture (TN)",   "Potential institutional partner; data consumer and policy stakeholder."),
]
add_table(["Stakeholder", "Role / Interest"], stakeholders, col_widths=[5.5, 10])


# ═══════════════════════════════════════════════════════════════════════════
#  4. FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════
h1("4. Functional Requirements")

# 4.1 Auth
h2("4.1  Authentication Module")
body("Routes: /login, /signup, /profile. All other routes are protected via ProtectedRoute.")
reqs = [
    ("FR-AUTH-01", "User can register with name, email, and password."),
    ("FR-AUTH-02", "User can log in with email and password."),
    ("FR-AUTH-03", "Authenticated session is maintained across page navigation."),
    ("FR-AUTH-04", "Unauthenticated access to protected routes redirects to /login."),
    ("FR-AUTH-05", "User can view and edit profile details at /profile."),
    ("FR-AUTH-06", "User can log out and session is cleared."),
]
add_table(["ID", "Requirement"], reqs, col_widths=[3.5, 12])

# 4.2 Home
h2("4.2  Home Dashboard")
body("Route: /  —  Entry point after login.")
reqs = [
    ("FR-HOME-01", "Display VigorTerra branding, tagline, and feature chips."),
    ("FR-HOME-02", "Show three featured agricultural insights (climate, soil, data-driven decisions)."),
    ("FR-HOME-03", "List six curated external resources (Ministry, ONAGRI, Agridata, FAO, etc.) as clickable cards."),
    ("FR-HOME-04", "All content adapts to selected language (EN / FR / AR)."),
]
add_table(["ID", "Requirement"], reqs, col_widths=[3.5, 12])

# 4.3 Sources
h2("4.3  Data Sources Page")
body("Route: /sources")
reqs = [
    ("FR-SRC-01", "Display the list of data providers used to build the dataset."),
    ("FR-SRC-02", "Each source shows name, type, URL reference, and coverage period."),
    ("FR-SRC-03", "Page is bilingual (EN/FR/AR)."),
]
add_table(["ID", "Requirement"], reqs, col_widths=[3.5, 12])

# 4.4 Dataset
h2("4.4  Dataset Overview Page")
body("Route: /dataset")
reqs = [
    ("FR-DATA-01", "Display dataset summary statistics (rows, columns, date range, key variables)."),
    ("FR-DATA-02", "Show per-variable metadata: name, type, unit, description."),
    ("FR-DATA-03", "Highlight missing-value counts and data quality notes."),
]
add_table(["ID", "Requirement"], reqs, col_widths=[3.5, 12])

# 4.5 Pipeline
h2("4.5  ML Pipeline Page")
body("Route: /pipeline")
reqs = [
    ("FR-PIPE-01", "Visualize the end-to-end ML pipeline (data collection → preprocessing → training → deployment)."),
    ("FR-PIPE-02", "Each pipeline step displays name, description, and associated tools/technologies."),
    ("FR-PIPE-03", "Pipeline is rendered responsively for mobile and desktop."),
]
add_table(["ID", "Requirement"], reqs, col_widths=[3.5, 12])

# 4.6 Prediction
h2("4.6  Yield Prediction (Test) Page")
body("Route: /test  —  Core functional page.")
reqs = [
    ("FR-PRED-01", "User selects one of three ML models: Random Forest, Gradient Boosting, or Stacking."),
    ("FR-PRED-02", "User enters 8 input features: rainfall (mm), temperature (°C), humidity (%), pH, nitrogen (%), phosphorus (mg/kg), potassium (mg/kg), cultivated area (ha)."),
    ("FR-PRED-03", "Validated form submission calls POST /api/predict."),
    ("FR-PRED-04", "Response displays predicted yield in t/ha with quality label (low / moderate / good / excellent)."),
    ("FR-PRED-05", "Anomaly flag and reason are displayed when any parameter is out of agronomic range."),
    ("FR-PRED-06", "User can trigger an AI explanation of the result via the agent endpoint."),
    ("FR-PRED-07", "Input validation prevents submission with missing or out-of-bound values."),
    ("FR-PRED-08", "Loading state is shown while API request is in flight."),
]
add_table(["ID", "Requirement"], reqs, col_widths=[3.5, 12])

# 4.7 AI Agent
h2("4.7  AI Agricultural Assistant — Ask Me")
body("Route: /askme  —  Conversational interface powered by Terra.")
reqs = [
    ("FR-AI-01", "User types a free-text farming question and submits."),
    ("FR-AI-02", "System first resolves the question against a local knowledge base (instant, no API)."),
    ("FR-AI-03", "If no local match, system calls Gemini 1.5 Flash via POST /api/ask."),
    ("FR-AI-04", "Conversation history is displayed in a chat-style interface."),
    ("FR-AI-05", "Rate limiting (15-second minimum interval) is enforced server-side for external API calls."),
    ("FR-AI-06", "Missing or invalid API key returns a clear error message to the user."),
    ("FR-AI-07", "User can request a plain-language explanation of a yield prediction result via POST /api/explain-yield."),
]
add_table(["ID", "Requirement"], reqs, col_widths=[3.5, 12])


# ═══════════════════════════════════════════════════════════════════════════
#  5. NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════
h1("5. Non-Functional Requirements")

nfr = [
    ("Performance",    "NFR-PERF-01",  "Prediction API response time < 500 ms under normal load."),
    ("Performance",    "NFR-PERF-02",  "Frontend initial load < 3 s on a 10 Mbps connection."),
    ("Availability",   "NFR-AVAIL-01", "Backend uptime ≥ 99% with restart policy ('unless-stopped')."),
    ("Scalability",    "NFR-SCAL-01",  "Docker-compose architecture enables horizontal backend scaling without frontend changes."),
    ("Security",       "NFR-SEC-01",   "CORS policy restricts origins to localhost:3000/3001/3002."),
    ("Security",       "NFR-SEC-02",   "Gemini API key is injected via environment variable; never hard-coded."),
    ("Security",       "NFR-SEC-03",   "All user-submitted inputs are validated via Pydantic field constraints before processing."),
    ("Usability",      "NFR-USE-01",   "UI supports three languages (EN, FR, AR) with dynamic RTL layout for Arabic."),
    ("Usability",      "NFR-USE-02",   "All interactive elements have visible loading and error states."),
    ("Maintainability","NFR-MAINT-01", "Backend router modules are independent and individually testable."),
    ("Maintainability","NFR-MAINT-02", "Frontend pages follow a consistent Layout wrapper component pattern."),
    ("Portability",    "NFR-PORT-01",  "Full stack runs on any Docker-enabled host (Linux, macOS, Windows)."),
]
add_table(["Category", "ID", "Requirement"], nfr, col_widths=[3.5, 3.5, 9])


# ═══════════════════════════════════════════════════════════════════════════
#  6. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
h1("6. System Architecture")

h2("6.1  Frontend Architecture")
body("Technology: React 18, React Router v6, Vite 5.")
body("Pattern: Single-Page Application (SPA) with client-side routing.")
components = [
    ("App.jsx",           "Root router — defines all route/page mappings and ProtectedRoute guards."),
    ("context/AuthContext.jsx",     "Global authentication state (login status, user object, login/logout actions)."),
    ("context/LanguageContext.jsx", "Global language state — drives all multilingual content switches."),
    ("components/Layout.jsx",       "Shared page wrapper: Header + Footer + main content slot."),
    ("components/Header.jsx",       "Navigation bar with links to all protected pages."),
    ("components/LanguageSwitcher.jsx", "Floating widget for EN / FR / AR selection."),
    ("pages/Login.jsx + SignUp.jsx","Authentication forms."),
    ("pages/Home.jsx",              "Dashboard with insights and resource links."),
    ("pages/TestPage.jsx",          "Yield prediction form and result display."),
    ("pages/AskMePage.jsx",         "Chat interface for Terra AI assistant."),
    ("pages/PipelinePage.jsx",      "ML pipeline visualization."),
    ("pages/DatasetPage.jsx",       "Dataset statistics and metadata table."),
    ("pages/SourcesPage.jsx",       "Data source catalog."),
    ("pages/Profile.jsx",           "User profile view and edit."),
]
add_table(["File / Module", "Responsibility"], components, col_widths=[5.5, 10])

h2("6.2  Backend Architecture")
body("Technology: FastAPI 0.115, Uvicorn, Pydantic v2, Python 3.12+.")
body("Pattern: Router-based modular API — each domain has its own router file.")
modules = [
    ("app/main.py",            "FastAPI application factory — mounts routers, configures CORS middleware."),
    ("app/routers/health.py",  "GET /health and GET /api/health — service liveness checks."),
    ("app/routers/predict.py", "POST /api/predict — yield prediction with model selection, quality labelling, anomaly detection."),
    ("app/routers/agent.py",   "POST /api/ask — AI chat (local KB + Gemini fallback). POST /api/explain-yield — result explanation."),
]
add_table(["Module", "Responsibility"], modules, col_widths=[4.5, 11])

h2("6.3  Data Flow")
body("Standard prediction flow:")
steps = [
    "User fills the prediction form in TestPage.jsx and selects a model.",
    "Frontend sends POST /api/predict with JSON payload (model + 8 features).",
    "FastAPI validates input via Pydantic (PredictionRequest model).",
    "Selected model predictor function computes yield (t/ha).",
    "Quality label and anomaly check are applied to the result.",
    "PredictionResponse JSON is returned to the frontend.",
    "Frontend renders yield value, quality badge, and anomaly alerts.",
    "(Optional) User clicks 'Explain' → POST /api/explain-yield → Terra explanation rendered in chat.",
]
for i, s in enumerate(steps, 1):
    bullet(f"Step {i}: {s}")


# ═══════════════════════════════════════════════════════════════════════════
#  7. API SPECIFICATION
# ═══════════════════════════════════════════════════════════════════════════
h1("7. API Specification")
body("Base URL (development): http://localhost:8000   |   Interactive docs: http://localhost:8000/docs")

endpoints = [
    ("GET",  "/health",           "—",                        "Service liveness check. Returns {status: ok}."),
    ("GET",  "/api/health",       "—",                        "Same check accessible via Vite /api proxy."),
    ("POST", "/api/predict",      "PredictionRequest (JSON)",  "Yield prediction. Returns PredictionResponse."),
    ("POST", "/api/ask",          "AskRequest {question}",     "AI chat. Returns AskResponse {answer}."),
    ("POST", "/api/explain-yield","ExplainYieldRequest (JSON)","Plain-language yield result explanation."),
]
add_table(["Method", "Path", "Request Body", "Description"], endpoints, col_widths=[2, 4.5, 4.5, 5])

h2("7.1  PredictionRequest Schema")
fields = [
    ("model",         "string",  "gradient_boosting",  "One of: random_forest, gradient_boosting, stacking"),
    ("pluviometrie",  "float",   "≥ 0",                "Rainfall in mm"),
    ("temperature",   "float",   "any",                "Average temperature in °C"),
    ("humidite",      "float",   "0–100",              "Relative humidity in %"),
    ("ph",            "float",   "0–14",               "Soil pH"),
    ("azote",         "float",   "≥ 0",                "Nitrogen content (%)"),
    ("phosphore",     "float",   "≥ 0",                "Phosphorus content (mg/kg)"),
    ("potassium",     "float",   "≥ 0",                "Potassium content (mg/kg)"),
    ("surface",       "float",   "≥ 0",                "Cultivated area in hectares"),
]
add_table(["Field", "Type", "Constraints", "Description"], fields, col_widths=[3, 2, 2.5, 8.5])

h2("7.2  PredictionResponse Schema")
resp_fields = [
    ("rendement_t_ha",   "float",  "Predicted yield in t/ha"),
    ("unit",             "string", "Always 't/ha'"),
    ("model",            "string", "Model used for this prediction"),
    ("rendement_quality","string", "low | moderate | good | excellent"),
    ("anomaly_detected", "bool",   "True if any input is outside agronomic range"),
    ("anomaly_reason",   "string?","Semicolon-separated list of anomaly descriptions (null if none)"),
]
add_table(["Field", "Type", "Description"], resp_fields, col_widths=[4, 2.5, 9.5])


# ═══════════════════════════════════════════════════════════════════════════
#  8. MACHINE LEARNING MODELS
# ═══════════════════════════════════════════════════════════════════════════
h1("8. Machine Learning Models")

body(
    "The current backend implements three deterministic approximation functions that simulate "
    "ML model behaviour. These are designed to mirror the behaviour of trained models without "
    "requiring serialised model files, making the API self-contained and instantly deployable."
)

h2("8.1  Random Forest Approximation")
body("File: app/routers/predict.py — _compute_prediction_random_forest()")
factors = [
    ("Base yield",    "2.5 t/ha constant baseline"),
    ("Rainfall",      "Linear capped effect: min(rainfall/400, 1.5) × 0.8"),
    ("Temperature",   "+0.3 if 15–25°C optimal range, else +0.1"),
    ("Soil pH",       "+0.4 if 6–8 optimal range, else +0.1"),
    ("NPK nutrients", "(N×10 + P/50 + K/200) × 0.1"),
    ("Tuning offset", "+0.08"),
]
add_table(["Factor", "Formula / Logic"], factors, col_widths=[4, 12])

h2("8.2  Gradient Boosting Approximation")
body("File: app/routers/predict.py — _compute_prediction_gradient_boosting()")
body("Builds on Random Forest with a non-linear correction term:")
body("  yield = RF_yield + 0.18 + (phosphorus × 0.0012) + (temperature − 22) × 0.01", indent=1)

h2("8.3  Stacking Approximation")
body("File: app/routers/predict.py — _compute_prediction_stacking()")
body("Meta-learner combination of RF and GB predictions:")
body("  yield = 0.3 × RF_yield + 0.7 × GB_yield", indent=1)

h2("8.4  Quality Labels")
quality = [
    ("low",      "< 2.5 t/ha",  "Below optimal. Significant improvement possible."),
    ("moderate", "2.5–4.0 t/ha","Average production. Review limiting factors."),
    ("good",     "4.0–6.0 t/ha","Solid yield. Minor optimisations may help."),
    ("excellent", "≥ 6.0 t/ha", "High production level. Conditions are favourable."),
]
add_table(["Label", "Range", "Interpretation"], quality, col_widths=[3, 3.5, 9.5])

h2("8.5  Anomaly Detection Rules")
anomalies = [
    ("Soil pH",         "Outside 5.5–8.5",    "pH out of agronomic range"),
    ("Temperature",     "< 5°C or > 40°C",    "temperature out of expected range"),
    ("Humidity",        "< 20% or > 95%",      "humidity out of expected range"),
    ("Rainfall",        "< 80 mm or > 1200 mm","rainfall out of expected range"),
    ("Predicted yield", "< 0.5 or > 15 t/ha", "predicted yield outlier"),
]
add_table(["Parameter", "Trigger Condition", "Reason String"], anomalies, col_widths=[3.5, 4.5, 8])

h2("8.6  Jupyter / MLflow Notebooks")
body("Full model training experiments are maintained in MLflow/:")
for nb in [
    "Yield_Prediction.ipynb — regression model for yield in t/ha.",
    "Productivity_Classification.ipynb — crop productivity class prediction.",
    "Disease_Risk.ipynb — disease risk index estimation.",
]:
    bullet(nb)


# ═══════════════════════════════════════════════════════════════════════════
#  9. DATA SOURCES
# ═══════════════════════════════════════════════════════════════════════════
h1("9. Data Sources")
sources = [
    ("Agridata Tunisia",     "agridata.tn",                  "Agricultural production and pluviometry statistics for Tunisia"),
    ("FAOSTAT",              "fao.org/faostat",              "Global crop production data, 1980–2024"),
    ("Open-Meteo",           "open-meteo.com",               "Historical climate data (temperature, rainfall) for Tunisian regions"),
    ("Soil data (generated)","agriculture_scraping/scripts", "Synthetic soil properties derived from regional agronomic profiles"),
    ("Ministry of Agric. TN","agriculture.tn",               "Official sector reports and technical factsheets"),
    ("ONAGRI",               "onagri.nat.tn",                "National agricultural observatory statistics"),
]
add_table(["Source", "Origin", "Content"], sources, col_widths=[4, 4, 8])
body("Scripts in agriculture_scraping/scripts/ automate collection and merge all sources into a single dataset at DATA/dataset_filtré.csv.")


# ═══════════════════════════════════════════════════════════════════════════
#  10. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════
h1("10. Technology Stack")
stack = [
    ("Frontend Framework",    "React 18.2",           "UI component library"),
    ("Frontend Build Tool",   "Vite 5.0",             "Development server and production bundler"),
    ("Frontend Routing",      "React Router v6",      "Client-side page routing with protected routes"),
    ("Backend Framework",     "FastAPI 0.115",        "High-performance Python REST API"),
    ("ASGI Server",           "Uvicorn 0.32",         "Production-grade async server"),
    ("Data Validation",       "Pydantic v2",          "Request/response schema validation"),
    ("AI Integration",        "Google Gemini 1.5 Flash","Generative AI for the Terra chat assistant"),
    ("Env Management",        "python-dotenv",        "Environment variable loading from .env"),
    ("Data Collection",       "BeautifulSoup4, lxml", "Web scraping for agridata.tn"),
    ("Data Processing",       "pandas, openpyxl",     "Dataset cleaning, merging, feature engineering"),
    ("ML Experiments",        "MLflow + Jupyter",     "Model training, tracking, and evaluation"),
    ("Containerisation",      "Docker + docker-compose","Service isolation and one-command deployment"),
    ("HTTP Proxy (dev)",      "Vite proxy config",    "Forwards /api/* from port 3000 to port 8000"),
    ("CSS Methodology",       "Component-scoped CSS", "Each component has its own .css file"),
]
add_table(["Layer", "Technology", "Purpose"], stack, col_widths=[4, 4, 8])


# ═══════════════════════════════════════════════════════════════════════════
#  11. DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════
h1("11. Deployment")

h2("11.1  Local Development (Recommended for Development)")
steps_dev = [
    "cd C:\\Users\\LENOVO\\Desktop\\VigorTerra-master",
    "# Terminal 1 — Backend",
    "pip install -r app/requirements.txt",
    "$env:GEMINI_API_KEY = 'your_key'",
    "uvicorn app.main:app --reload --host 0.0.0.0 --port 8000",
    "",
    "# Terminal 2 — Frontend",
    "npm install",
    "npm run dev   # starts on http://localhost:3000",
]
for step in steps_dev:
    if step and not step.startswith("#"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(step)
        set_font(r, name="Courier New", size=9)
    elif step.startswith("#"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(step)
        set_font(r, name="Courier New", size=9, color=(100, 100, 100), italic=True)

h2("11.2  Docker Deployment (Production)")
body("Copy .env.example to .env and set GEMINI_API_KEY, then:")
for cmd in [
    "docker-compose up --build",
    "# Frontend → http://localhost:80",
    "# Backend  → http://localhost:8000",
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(cmd)
    color = (100, 100, 100) if cmd.startswith("#") else None
    set_font(r, name="Courier New", size=9, italic=cmd.startswith("#"), color=color)

h2("11.3  Environment Variables")
env_vars = [
    ("GEMINI_API_KEY", "Required", "Google Gemini API key for the Terra AI assistant. Obtain free at aistudio.google.com."),
    ("VITE_API_URL",   "Build-time (Docker)", "Backend URL injected at build time for the Docker frontend image. Default: http://localhost:8000."),
]
add_table(["Variable", "Status", "Description"], env_vars, col_widths=[4, 3, 9])

h2("11.4  Port Summary")
ports = [
    ("3000", "Vite dev server (frontend)", "Development only"),
    ("5173", "Vite default fallback port", "If port 3000 is taken"),
    ("8000", "Uvicorn FastAPI backend",    "Development and Docker"),
    ("80",   "Nginx (Docker frontend)",   "Production Docker only"),
]
add_table(["Port", "Service", "Environment"], ports, col_widths=[2, 7, 7])


# ═══════════════════════════════════════════════════════════════════════════
#  12. CONSTRAINTS AND ASSUMPTIONS
# ═══════════════════════════════════════════════════════════════════════════
h1("12. Constraints and Assumptions")

h2("12.1  Technical Constraints")
for c in [
    "ML models are rule-based approximations, not trained on real labelled Tunisian yield data. Accuracy is indicative, not certified.",
    "The Gemini AI integration is subject to Google's free-tier rate limits (15-second minimum between API calls enforced server-side).",
    "Authentication is implemented via React Context (client-side); there is no persistent user database in the current version.",
    "CORS is restricted to localhost origins — production deployment requires CORS configuration update.",
    "Python 3.12+ is required for the backend due to PEP 604 union type hints (X | Y syntax).",
]:
    bullet(c)

h2("12.2  Assumptions")
for a in [
    "Users have access to basic soil and climate data for their fields.",
    "The platform is used in a desktop/laptop browser environment.",
    "Internet access is available for Gemini API calls (local knowledge base works offline).",
    "Tunisian agronomic ranges are used as anomaly detection thresholds.",
]:
    bullet(a)


# ═══════════════════════════════════════════════════════════════════════════
#  13. GLOSSARY
# ═══════════════════════════════════════════════════════════════════════════
h1("13. Glossary")
terms = [
    ("t/ha",               "Tonnes per hectare — standard unit for crop yield measurement."),
    ("NPK",                "Nitrogen (N), Phosphorus (P), Potassium (K) — the three primary soil nutrients."),
    ("pH",                 "Power of Hydrogen — scale (0–14) measuring soil acidity/alkalinity."),
    ("ML",                 "Machine Learning — computational models that learn patterns from data."),
    ("Random Forest",      "Ensemble ML model using multiple decision trees and majority voting."),
    ("Gradient Boosting",  "Ensemble ML model that builds trees sequentially to correct prior errors."),
    ("Stacking",           "Meta-ensemble that combines predictions from multiple base models."),
    ("Anomaly",            "A prediction or input value that falls outside expected agronomic ranges."),
    ("Terra",              "The AI agricultural assistant integrated into VigorTerra's Ask Me page."),
    ("FAOSTAT",            "Food and Agriculture Organization Statistical Database — global agricultural data."),
    ("ONAGRI",             "Observatoire National de l'Agriculture — Tunisia's national agricultural observatory."),
    ("FastAPI",            "Modern Python web framework for building REST APIs with automatic OpenAPI documentation."),
    ("Vite",               "Next-generation frontend build tool and development server for React/Vue/etc."),
    ("Pydantic",           "Python data validation library used for API request/response schemas."),
    ("Docker",             "Container platform for packaging and deploying applications consistently."),
    ("Cahier de Charge",   "French term for 'specifications document' — defines functional and technical requirements."),
]
add_table(["Term", "Definition"], terms, col_widths=[4, 12])


# ── Footer / document end ────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"VigorTerra — Software Specifications v1.0  |  Generated {datetime.date.today().strftime('%B %d, %Y')}  |  CONFIDENTIAL")
set_font(run, size=8, italic=True, color=(140, 140, 140))

# ── Save ─────────────────────────────────────────────────────────────────
output_path = r"c:\Users\LENOVO\Desktop\VigorTerra-master\specifications.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
